"""
JD Analytics + Resume Parsing.

Pipeline:
  1. read raw text from a PDF or DOCX (pdfplumber / python-docx)
  2. hand the text to the configured LLM with a prompt that maps requirements
     onto the 12 RADIX categories -> structured skill list
  3. if no LLM is available, fall back to a keyword/heuristic extractor so the
     tool still produces *something* offline.
"""

from __future__ import annotations

import io
import os
import re

from . import llm
from .models import ExtractedSkillList, Skill
from .taxonomy import CATEGORIES, KEYWORDS, SKILLSET_CODES, coerce_code

# ---------------------------------------------------------------------------
# 1. text extraction
# ---------------------------------------------------------------------------

def extract_text(filename: str, data: bytes) -> str:
    """Return raw text from a PDF or DOCX given its bytes."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _pdf_text(data)
    if name.endswith(".docx"):
        return _docx_text(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    # last resort: try pdf then docx
    try:
        return _pdf_text(data)
    except Exception:
        return _docx_text(data)


def _pdf_text(data: bytes) -> str:
    import pdfplumber

    out: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n".join(out).strip()


def _docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts).strip()


# ---------------------------------------------------------------------------
# 2. LLM prompts
# ---------------------------------------------------------------------------

def _category_block() -> str:
    return "\n".join(f"  {code} = {name}" for code, name in CATEGORIES.items())


_COMMON_RULES = f"""You map technical hiring documents onto RADIX's 12 skill categories.

CATEGORY CODES (use exactly these codes):
{_category_block()}

Rules:
- Every skill you output MUST have a category_code from the list above. Use OTHER only when nothing fits.
- Prefer specific named technologies (e.g. "React", "PostgreSQL", "Kubernetes") AND the underlying skill
  (e.g. "System design at scale", "Data structures & algorithms") — capture both when present.
- evidence = a short quote or reason from the document (<= 12 words).
- confidence = "high" | "medium" | "low" based on how strongly the document signals the skill.
- Do NOT invent skills that are not supported by the text.
- Return ONLY a JSON object. No markdown, no commentary."""


def analyze_jd(text: str, source_file: str = "") -> ExtractedSkillList:
    system = _COMMON_RULES
    user = f"""Analyse this JOB DESCRIPTION. The strongest skill signal is usually in
"Key Responsibilities" and "What We're Looking For".

Return JSON of shape:
{{
  "company": "<hiring company>",
  "role": "<job title>",
  "skills": [{{"skill_name": "...", "category_code": "...", "evidence": "...", "confidence": "high|medium|low"}}]
}}

JOB DESCRIPTION:
\"\"\"
{text[:12000]}
\"\"\""""

    data = llm.complete_json(system, user)
    if data is None:
        return _local_extract(text, "jd", source_file)

    return ExtractedSkillList(
        source_type="jd",
        source_file=source_file,
        company=str(data.get("company", "")),
        role=str(data.get("role", "")),
        skills=_coerce_skills(data.get("skills", [])),
    )


def analyze_resume(text: str, source_file: str = "") -> ExtractedSkillList:
    system = _COMMON_RULES
    user = f"""Analyse this RESUME. Capture skills from the skills section AND skills
implied by projects, internships, and achievements.

Return JSON of shape:
{{
  "name": "<candidate name>",
  "education": "<one-line education summary>",
  "projects": ["<short project descriptions>"],
  "experience": ["<internships / work>"],
  "role": "<preferred role if stated>",
  "skills": [{{"skill_name": "...", "category_code": "...", "evidence": "...", "confidence": "high|medium|low"}}]
}}

RESUME:
\"\"\"
{text[:12000]}
\"\"\""""

    data = llm.complete_json(system, user)
    if data is None:
        result = _local_extract(text, "resume", source_file)
        result.education = _guess_education(text)
        return result

    return ExtractedSkillList(
        source_type="resume",
        source_file=source_file,
        name=str(data.get("name", "")),
        role=str(data.get("role", "")),
        skills=_coerce_skills(data.get("skills", [])),
        education=str(data.get("education", "")),
        projects=[str(p) for p in data.get("projects", []) if p],
        experience=[str(e) for e in data.get("experience", []) if e],
    )


def _coerce_skills(raw: list) -> list[Skill]:
    skills: list[Skill] = []
    seen: set[tuple[str, str]] = set()
    for item in raw or []:
        if not isinstance(item, dict):
            continue
        name = str(item.get("skill_name", "")).strip()
        if not name:
            continue
        code = coerce_code(item.get("category_code"))
        key = (name.lower(), code)
        if key in seen:
            continue
        seen.add(key)
        conf = str(item.get("confidence", "medium")).lower()
        if conf not in ("high", "medium", "low"):
            conf = "medium"
        skills.append(
            Skill(
                skill_name=name,
                category_code=code,
                evidence=str(item.get("evidence", ""))[:160],
                confidence=conf,  # type: ignore[arg-type]
            )
        )
    return skills


# ---------------------------------------------------------------------------
# 3. offline fallback (no API key)
# ---------------------------------------------------------------------------

def _local_extract(text: str, source_type: str, source_file: str) -> ExtractedSkillList:
    lowered = f" {text.lower()} "
    skills: list[Skill] = []
    for code in SKILLSET_CODES:
        hits = [kw for kw in KEYWORDS[code] if kw in lowered]
        if not hits:
            continue
        conf = "high" if len(hits) >= 3 else "medium" if len(hits) == 2 else "low"
        skills.append(
            Skill(
                skill_name=CATEGORIES[code],
                category_code=code,
                evidence="keywords: " + ", ".join(h.strip() for h in hits[:4]),
                confidence=conf,  # type: ignore[arg-type]
            )
        )
    return ExtractedSkillList(
        source_type=source_type,  # type: ignore[arg-type]
        source_file=source_file,
        skills=skills,
    )


def _guess_education(text: str) -> str:
    m = re.search(r"(b\.?tech|m\.?tech|b\.?e\.?|b\.?sc|m\.?sc|bachelor|master).{0,80}", text, re.I)
    return m.group(0).strip() if m else ""
